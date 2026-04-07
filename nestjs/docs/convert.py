# -*- coding: utf-8 -*-
"""
将 doc（实际为 docx 格式）文件转换为 markdown 文件
提取图片到 nestjs/images 目录

使用方法：
  pip install python-docx lxml
  python convert.py
"""
import os
import json
from docx import Document

# 命名空间
nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
}

docs_dir = os.path.dirname(os.path.abspath(__file__))

print(docs_dir, "docs_dir")
nestjs_dir = os.path.dirname(docs_dir)
images_dir = os.path.join(nestjs_dir, 'images')
os.makedirs(images_dir, exist_ok=True)

for filename in sorted(os.listdir(docs_dir)):
    if not filename.endswith('.doc'):
        continue

    filepath = os.path.join(docs_dir, filename)
    doc = Document(filepath)

    # 提取文件基础名（去掉 _课件.doc 后缀）
    base_name = filename.replace('_课件.doc', '')
    md_filename = base_name + '.md'

    # 提取所有图片关系
    image_rels = {}
    img_counter = 0
    for rel_id, rel in doc.part.rels.items():
        if 'image' in rel.reltype:
            img_counter += 1
            img_data = rel.target_part.blob
            content_type = rel.target_part.content_type
            ext = '.png'
            if 'jpeg' in content_type or 'jpg' in content_type:
                ext = '.jpg'
            elif 'gif' in content_type:
                ext = '.gif'

            img_filename = f'{base_name}_img{img_counter}{ext}'
            img_path = os.path.join(images_dir, img_filename)
            with open(img_path, 'wb') as f:
                f.write(img_data)
            image_rels[rel_id] = img_filename

    # 按段落顺序生成 markdown
    md_lines = [f'# {base_name}\n']

    for para in doc.paragraphs:
        text = para.text.strip()

        # 检查段落中的图片
        blips = para._element.findall('.//a:blip', nsmap)
        for blip in blips:
            embed = blip.get(
                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            )
            if embed and embed in image_rels:
                img_name = image_rels[embed]
                md_lines.append(f'\n![{img_name}](./images/{img_name})\n')

        if not text:
            continue

        style_name = para.style.name if para.style else ''

        # 根据样式转换标题层级
        if 'Heading 1' in style_name or style_name == '标题 1':
            md_lines.append(f'\n## {text}\n')
        elif 'Heading 2' in style_name or style_name == '标题 2':
            md_lines.append(f'\n### {text}\n')
        elif 'Heading 3' in style_name or style_name == '标题 3':
            md_lines.append(f'\n#### {text}\n')
        else:
            md_lines.append(f'{text}\n')

    md_content = '\n'.join(md_lines)
    md_path = os.path.join(nestjs_dir, md_filename)
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    print(f'已生成: {md_filename} (图片: {img_counter} 张, 段落: {len(doc.paragraphs)})')

print('\n所有文件转换完成!')
